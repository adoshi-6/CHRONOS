"""
CHRONOS Universal File Upload Readers (file_dispatcher.py)
------------------------------------------------------------
Multi-format file parser for CHRONOS supporting PDF, Word (.docx),
Excel (.xlsx), Images (.png/.jpg), Text, CSV, JSON, DXF CAD, and OpenSCAD 3D files.
"""

import os
import json
import traceback

class UniversalFileDispatcher:
    """Dispatches files to specialized format parsers based on extension."""

    def __init__(self):
        self.supported_extensions = {
            "text": [".txt", ".md", ".json", ".csv", ".py", ".js", ".html", ".css", ".yaml", ".yml", ".xml", ".log", ".sql"],
            "pdf": [".pdf"],
            "word": [".docx", ".doc"],
            "excel": [".xlsx", ".xls"],
            "image": [".png", ".jpg", ".jpeg", ".webp", ".bmp", ".gif"],
            "cad_3d": [".dxf", ".scad", ".obj", ".stl"]
        }

    def parse_file(self, filepath):
        """Main dispatcher entrypoint. Parses file and returns structured output."""
        if not os.path.exists(filepath):
            return {"success": False, "error": f"File not found: {filepath}"}

        filename = os.path.basename(filepath)
        ext = os.path.splitext(filename)[1].lower()
        filesize_bytes = os.path.getsize(filepath)

        result = {
            "filename": filename,
            "filepath": filepath,
            "extension": ext,
            "size_bytes": filesize_bytes,
            "success": True,
            "content": "",
            "metadata": {}
        }

        try:
            if ext in self.supported_extensions["text"]:
                result["type"] = "Text/Code"
                result["content"] = self._parse_text(filepath)
            elif ext in self.supported_extensions["pdf"]:
                result["type"] = "PDF Document"
                result["content"], result["metadata"] = self._parse_pdf(filepath)
            elif ext in self.supported_extensions["word"]:
                result["type"] = "Word Document"
                result["content"], result["metadata"] = self._parse_word(filepath)
            elif ext in self.supported_extensions["excel"]:
                result["type"] = "Excel Spreadsheet"
                result["content"], result["metadata"] = self._parse_excel(filepath)
            elif ext in self.supported_extensions["image"]:
                result["type"] = "Image Media"
                result["content"], result["metadata"] = self._parse_image(filepath)
            elif ext in self.supported_extensions["cad_3d"]:
                result["type"] = "CAD / 3D Geometry"
                result["content"], result["metadata"] = self._parse_cad_3d(filepath)
            else:
                result["type"] = "Binary/Generic"
                result["content"] = f"[Binary file format '{ext}' - Raw content skipped]"
        except Exception as e:
            result["success"] = False
            result["error"] = f"Failed to parse {filename}: {str(e)}"
            result["traceback"] = traceback.format_exc()

        return result

    def _parse_text(self, filepath):
        """Parses plain text and code files with UTF-8 encoding."""
        with open(filepath, "r", encoding="utf-8", errors="replace") as f:
            return f.read()

    def _parse_pdf(self, filepath):
        """Parses PDF documents using PyPDF2 / pypdf or fallback byte stream."""
        content = ""
        metadata = {}
        try:
            import pypdf
            reader = pypdf.PdfReader(filepath)
            metadata["total_pages"] = len(reader.pages)
            pages_text = []
            for idx, page in enumerate(reader.pages):
                txt = page.extract_text() or ""
                pages_text.append(f"--- Page {idx + 1} ---\n{txt}")
            content = "\n".join(pages_text)
        except ImportError:
            try:
                import PyPDF2
                reader = PyPDF2.PdfReader(filepath)
                metadata["total_pages"] = len(reader.pages)
                content = "\n".join([page.extract_text() for page in reader.pages])
            except Exception:
                content = "[PDF parsing library (pypdf/PyPDF2) not installed. Please install via pip]."
                metadata["warning"] = "Missing PDF library"
        return content, metadata

    def _parse_word(self, filepath):
        """Parses Word (.docx) documents."""
        content = ""
        metadata = {}
        try:
            import docx
            doc = docx.Document(filepath)
            metadata["paragraphs_count"] = len(doc.paragraphs)
            metadata["tables_count"] = len(doc.tables)
            paras = [p.text for p in doc.paragraphs if p.text.strip()]
            content = "\n".join(paras)
        except ImportError:
            content = "[docx parsing library (python-docx) not installed]."
            metadata["warning"] = "Missing python-docx"
        return content, metadata

    def _parse_excel(self, filepath):
        """Parses Excel (.xlsx / .xls) spreadsheets."""
        content = ""
        metadata = {}
        try:
            import openpyxl
            wb = openpyxl.load_workbook(filepath, data_only=True)
            metadata["sheets"] = wb.sheetnames
            sheets_data = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                rows = []
                for row in ws.iter_rows(values_only=True):
                    if any(row):
                        rows.append(" | ".join([str(cell) if cell is not None else "" for cell in row]))
                sheets_data.append(f"=== Sheet: {sheet_name} ===\n" + "\n".join(rows[:100]))
            content = "\n\n".join(sheets_data)
        except ImportError:
            content = "[Excel parsing library (openpyxl) not installed]."
            metadata["warning"] = "Missing openpyxl"
        return content, metadata

    def _parse_image(self, filepath):
        """Parses image metadata using PIL/Pillow."""
        content = ""
        metadata = {}
        try:
            from PIL import Image
            with Image.open(filepath) as img:
                metadata["format"] = img.format
                metadata["size"] = img.size # (width, height)
                metadata["mode"] = img.mode
                content = f"Image Media [{img.format}] Resolution: {img.size[0]}x{img.size[1]}px, Color Mode: {img.mode}"
        except Exception:
            content = f"Image file: {os.path.basename(filepath)}"
        return content, metadata

    def _parse_cad_3d(self, filepath):
        """Parses DXF and OpenSCAD files."""
        content = ""
        metadata = {}
        with open(filepath, "r", encoding="utf-8", errors="replace") as f:
            raw_text = f.read()

        ext = os.path.splitext(filepath)[1].lower()
        if ext == ".scad":
            metadata["cad_type"] = "OpenSCAD Code"
            content = f"// OpenSCAD 3D Geometry Script\n{raw_text}"
        elif ext == ".dxf":
            metadata["cad_type"] = "AutoCAD DXF"
            content = f"// DXF Wireframe Entities\n{raw_text[:2000]}\n... [Content Truncated]"
        else:
            content = raw_text[:2000]

        return content, metadata


# Singleton instance
file_dispatcher = UniversalFileDispatcher()


if __name__ == "__main__":
    print("\n--- CHRONOS Universal File Dispatcher Diagnostics ---")
    demo_file = os.path.abspath(__file__)
    res = file_dispatcher.parse_file(demo_file)
    print(f"Parsed Self ({res['filename']}): {res['type']} - {res['size_bytes']} bytes")
    print("Content Preview:\n", res['content'][:200])
